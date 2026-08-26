"""
Extraktoren pro Dateityp. Jeder Extraktor nimmt einen Path und gibt (text, language) zurück.
Bei nicht unterstützten Typen: raise UnsupportedFormat.
Bei Fehler in der Extraktion: raise ExtractionError.

SICHERHEIT: Alle Extraktoren öffnen Dateien ausschliesslich lesend.
Verboten: os.remove, os.rename, shutil.delete, open(..., 'w'), open(..., 'wb').
"""

import logging
import multiprocessing
import re
import signal as _signal
import threading
from pathlib import Path

log = logging.getLogger(__name__)

_IN_WORKER_PROCESS = False  # wird von _scan_file_worker auf True gesetzt

_PDF_TIMEOUT = 120   # Sekunden für PyMuPDF / pypdf (Daemon-Thread, nur ausserhalb Worker)
_OCR_TIMEOUT = 7200  # Sicherheitsnetz für OCR-Prozess — per-Seiten-Timeout greift früher
_PAGE_TEXT_TIMEOUT = 15   # Sekunden pro Seite für Text-Extraktion (nur im Worker via SIGALRM)
_PAGE_OCR_TIMEOUT  = 60   # Sekunden pro Seite für OCR (nur im _worker_ocr-Prozess)
_NON_PDF_TIMEOUT   = 30   # Sekunden für alle Nicht-PDF-Formate (DOCX, EML, RTF, …)


class _NonPdfTimeout(Exception):
    pass


def _non_pdf_alarm(signum, frame):
    raise _NonPdfTimeout()


class _PageTimeout(Exception):
    pass


def _alarm_raise(signum, frame):
    raise _PageTimeout()


# ── OCR-Worker (Subprocess — Tesseract kann GIL halten) ──────────────────────

def _ocr_pages(path_str: str) -> list[dict]:
    """OCR aller Seiten. Pro Seite _PAGE_OCR_TIMEOUT via SIGALRM (nur im Hauptthread
    des jeweiligen Prozesses gültig). Gibt [{page_number, content}] zurück.
    Läuft je nach Kontext inline im Worker-Prozess ODER im Subprozess (_worker_ocr)."""
    import signal

    class _OCRPageTimeout(Exception):
        pass

    def _ocr_alarm(signum, frame):
        raise _OCRPageTimeout()

    import fitz
    FLAGS = fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_DEHYPHENATE
    result = []
    old_handler = signal.signal(signal.SIGALRM, _ocr_alarm)
    try:
        with fitz.open(path_str) as doc:
            for i, page in enumerate(doc, start=1):
                text = ""
                signal.alarm(_PAGE_OCR_TIMEOUT)
                try:
                    for lang in ("deu+eng", "deu", "eng"):
                        try:
                            tp = page.get_textpage_ocr(
                                flags=FLAGS, language=lang, dpi=150, full=True,
                            )
                            text = page.get_text(textpage=tp, flags=FLAGS).strip()
                            if text:
                                break
                        except _OCRPageTimeout:
                            raise
                        except Exception:
                            continue
                except _OCRPageTimeout:
                    log.warning("OCR Seite %d Timeout (%ds) übersprungen: %s",
                                i, _PAGE_OCR_TIMEOUT, Path(path_str).name)
                    text = ""
                finally:
                    signal.alarm(0)
                if len(text) >= 50:
                    result.append({"page_number": i, "content": text})
    finally:
        signal.alarm(0)
        signal.signal(signal.SIGALRM, old_handler)
    return result


def _worker_ocr(path_str: str, result_q: multiprocessing.Queue) -> None:
    """Subprozess-Wrapper (Server-Kontext): OCR im eigenen Prozess damit Tesseract den
    Server-GIL nicht blockiert. Ergebnis über die Queue."""
    try:
        result_q.put(("ok", _ocr_pages(path_str)))
    except Exception as e:
        result_q.put(("err", str(e)))


class UnsupportedFormat(Exception):
    pass


class ExtractionError(Exception):
    pass


# ── Typographische Normalisierung ─────────────────────────────────────────────
# OCR-PDFs enthalten oft Unicode-Ligaturen (ﬂ ﬁ ﬀ ﬃ ﬄ) anstelle normaler
# Buchstaben. "Geschossﬂäche" (mit U+FB02) trifft kein LIKE '%geschossfläche%'.

_LIGATURE_TABLE = str.maketrans({
    'ﬀ': 'ff',   # ﬀ LATIN SMALL LIGATURE FF
    'ﬁ': 'fi',   # ﬁ LATIN SMALL LIGATURE FI
    'ﬂ': 'fl',   # ﬂ LATIN SMALL LIGATURE FL
    'ﬃ': 'ffi',  # ﬃ LATIN SMALL LIGATURE FFI
    'ﬄ': 'ffl',  # ﬄ LATIN SMALL LIGATURE FFL
    'ﬅ': 'st',   # ﬅ LATIN SMALL LIGATURE LONG S T
    'ﬆ': 'st',   # ﬆ LATIN SMALL LIGATURE ST
    'ʼ': "'",    # ʼ MODIFIER LETTER APOSTROPHE (häufig in OCR)
    '’': "'",    # ' RIGHT SINGLE QUOTATION MARK
})


# OCR-Leerzeichen nach Ligatur-Zeichenfolgen:
# "Geschossfl äche" → "Geschossfläche"  (ﬂ wurde zu fl, dann Leerzeichen)
# "Aufl age" → "Auflage"
# "Defi nition" → "Definition"
# NUR nach den Ligatur-Buchstaben-Paaren (fl, fi, ff, ffi, ffl) — nicht generell,
# sonst würden echte Wortlücken wie "ein neues" zu "einneues".
_OCR_LIGATURE_SPACE = re.compile(r'(fl|fi|ff|ffi|ffl) (?=[a-zäöüß])')

# Mehrfach-Leerzeichen die OCR manchmal erzeugt ("Defi  nition")
_OCR_MULTI_SPACE = re.compile(r' {2,}')


def normalize_text(text: str) -> str:
    """Ligaturen und OCR-Leerzeichen-Artefakte normalisieren. Umlaute bleiben erhalten.

    Fixes:
    - ﬂ → fl, ﬁ → fi, ﬀ → ff, ﬃ → ffi, ﬄ → ffl  (Unicode-Ligaturen)
    - "Geschossfl äche" → "Geschossfläche"  (Leerzeichen nach Ligatur-Zeichenpaar)
    - "Aufl age" → "Auflage"
    - "Defi  nition" → "Definition"
    """
    text = text.translate(_LIGATURE_TABLE)
    text = _OCR_MULTI_SPACE.sub(' ', text)
    text = _OCR_LIGATURE_SPACE.sub(r'\1', text)
    return text


# ── TXT / MD ──────────────────────────────────────────────────────────────────

def extract_txt(path: Path) -> tuple[str, str]:
    # READ-ONLY: NAS darf nie verändert werden
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=enc), ""
        except UnicodeDecodeError:
            continue
    raise ExtractionError(f"Cannot decode {path}")


# ── PDF ───────────────────────────────────────────────────────────────────────

def _shrink_fitz_store() -> None:
    """Leert den MuPDF-internen Speicher-Cache nach jeder PDF-Verarbeitung."""
    try:
        import fitz
        fitz.TOOLS.store_shrink(100)
    except Exception:
        pass


def _has_mojibake(pages: list[dict]) -> bool:
    if not pages:
        return False
    sample = "".join(p["content"] for p in pages[:5])
    if not sample:
        return False
    return sample.count("�") / len(sample) > 0.30


def _run_pdf_in_thread(fn, path: Path, timeout: int = _PDF_TIMEOUT) -> list[dict]:
    """Führt PDF-Extraktion in Daemon-Thread aus.
    Bei Timeout: Thread läuft als Daemon weiter (kein Blockieren), Datei wird übersprungen.
    PyMuPDF/pypdf geben den GIL für I/O frei → Server bleibt responsiv.
    """
    result: list = []
    error:  list = []

    def _run():
        try:
            result.append(fn(path))
        except Exception as e:
            error.append(e)

    t = threading.Thread(target=_run, daemon=True)
    t.start()
    t.join(timeout=timeout)
    if t.is_alive():
        log.warning("PDF Timeout (%ds), übersprungen: %s", timeout, path.name)
        raise ExtractionError(f"Timeout {timeout}s: {path.name}")
    if error:
        raise error[0]
    return result[0] if result else []


def _run_pdf(fn, path: Path, timeout: int) -> list[dict]:
    """Im Worker-Prozess: direkt ausführen (Prozess = Isolation).
    Sonst: Daemon-Thread mit Timeout (Server-Kontext)."""
    if _IN_WORKER_PROCESS:
        return fn(path)
    return _run_pdf_in_thread(fn, path, timeout)


def _run_ocr_in_process(path: Path) -> list[dict]:
    """OCR. Im Worker-Prozess (Pool-Worker sind daemon → dürfen KEINE Kindprozesse
    starten, sonst 'daemonic processes are not allowed to have children') direkt inline
    ausführen — der Worker ist selbst schon isoliert und hat den 120s-SIGKILL des
    Scanners als harte Grenze, plus _PAGE_OCR_TIMEOUT pro Seite. Sonst (Server-Thread)
    in eigenem Prozess mit hartem SIGKILL nach _OCR_TIMEOUT (Tesseract blockiert den GIL).
    """
    if _IN_WORKER_PROCESS:
        return _ocr_pages(str(path))
    ctx = multiprocessing.get_context("spawn")
    q: multiprocessing.Queue = ctx.Queue()
    p = ctx.Process(target=_worker_ocr, args=(str(path), q), daemon=True)
    p.start()
    p.join(timeout=_OCR_TIMEOUT)
    if p.is_alive():
        p.kill()
        p.join(timeout=10)
        log.warning("OCR Timeout (%ds), übersprungen: %s", _OCR_TIMEOUT, path.name)
        raise ExtractionError(f"OCR Timeout {_OCR_TIMEOUT}s: {path.name}")
    if q.empty():
        raise ExtractionError(f"OCR kein Ergebnis: {path.name}")
    status, payload = q.get_nowait()
    if status == "err":
        raise ExtractionError(payload)
    return payload


def extract_pdf(path: Path) -> tuple[str, str]:
    pages = extract_pdf_pages(path)
    return "\n".join(p["content"] for p in pages), ""


_PLAN_APPS = {"archicad", "vectorworks"}


def extract_pdf_metadata(path: Path) -> dict:
    """Liest Creator/Producer aus PDF-Metadaten. Erkennt ArchiCAD/Vectorworks → is_plan."""
    creator = producer = ""
    try:
        import fitz
        with fitz.open(str(path)) as doc:
            m = doc.metadata or {}
            creator  = str(m.get("creator",  "") or "")
            producer = str(m.get("producer", "") or "")
    except Exception:
        try:
            import pypdf
            info     = pypdf.PdfReader(str(path)).metadata or {}
            creator  = str(info.get("/Creator",  "") or "")
            producer = str(info.get("/Producer", "") or "")
        except Exception:
            return {}
    combined = (creator + " " + producer).lower()
    creator_app = next((a for a in _PLAN_APPS if a in combined), "")
    return {"creator_app": creator_app, "is_plan": bool(creator_app)}


def extract_pdf_pages(path: Path) -> list[dict]:
    """Gibt [{page_number, content}] zurück.

    Reihenfolge:
    1. PyMuPDF   — schnell, beste Qualität (Daemon-Thread, 120s Timeout)
    2. pypdf     — Fallback wenn PyMuPDF fehlt/scheitert (Daemon-Thread, 120s Timeout)
    3. OCR       — nur wenn 1+2 keinen lesbaren Text liefern (Subprocess, 90s SIGKILL)

    Jeder Schritt wird übersprungen wenn der vorherige lesbare Seiten liefert.
    Dateien die alle drei Schritte nicht lesen können: status='error', beim nächsten
    Scan übersprungen (solange Datei unverändert bleibt).
    """
    # ── 1. PyMuPDF ───────────────────────────────────────────────────────────
    try:
        def _pymupdf(p):
            import fitz
            FLAGS = fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_DEHYPHENATE
            pages = []
            # Per-Seiten-Timeout nur im Worker (Hauptthread → SIGALRM verfügbar).
            # Im Server-Thread würde signal.signal ValueError werfen.
            use_alarm = _IN_WORKER_PROCESS
            if use_alarm:
                old_handler = _signal.signal(_signal.SIGALRM, _alarm_raise)
            try:
                with fitz.open(str(p)) as doc:
                    for i, page in enumerate(doc, start=1):
                        if use_alarm:
                            _signal.alarm(_PAGE_TEXT_TIMEOUT)
                        try:
                            text = page.get_text(flags=FLAGS).strip()
                        except _PageTimeout:
                            log.warning("Seite %d Text-Timeout (%ds) übersprungen: %s",
                                        i, _PAGE_TEXT_TIMEOUT, Path(p).name)
                            text = ""
                        except Exception as exc:
                            log.debug("Seite %d Fehler: %s – %s", i, Path(p).name, exc)
                            text = ""
                        finally:
                            if use_alarm:
                                _signal.alarm(0)
                        if len(text) >= 50:
                            pages.append({"page_number": i, "content": text})
            finally:
                if use_alarm:
                    _signal.alarm(0)
                    _signal.signal(_signal.SIGALRM, old_handler)
            return pages

        pages = _run_pdf(_pymupdf, path, _PDF_TIMEOUT)
        if pages and not _has_mojibake(pages):
            _shrink_fitz_store()
            return pages
        if pages:
            log.info("PyMuPDF Mojibake erkannt: %s", path.name)
    except ImportError:
        log.debug("fitz nicht installiert: %s", path.name)
        pages = []
    except ExtractionError:
        log.info("PyMuPDF Timeout/Fehler: %s — versuche pypdf", path.name)
        pages = []
    except Exception as e:
        log.info("PyMuPDF Fehler (%s): %s — versuche pypdf", path.name, e)
        pages = []

    # ── 2. pypdf ─────────────────────────────────────────────────────────────
    try:
        def _pypdf(p):
            import pypdf
            result = []
            for i, page in enumerate(pypdf.PdfReader(str(p)).pages, start=1):
                text = (page.extract_text() or "").strip()
                if len(text) >= 50:
                    result.append({"page_number": i, "content": text})
            return result

        pypdf_pages = _run_pdf(_pypdf, path, _PDF_TIMEOUT)
        if pypdf_pages and not _has_mojibake(pypdf_pages):
            return pypdf_pages
        if pypdf_pages:
            log.info("pypdf Mojibake erkannt: %s", path.name)
    except ImportError:
        if not pages:
            raise UnsupportedFormat("Weder pymupdf noch pypdf installiert")
    except ExtractionError:
        log.info("pypdf Timeout/Fehler: %s — versuche OCR", path.name)
        pypdf_pages = []
    except Exception as e:
        log.info("pypdf Fehler (%s): %s — versuche OCR", path.name, e)
        pypdf_pages = []

    # ── 3. OCR (letzter Ausweg) ───────────────────────────────────────────────
    try:
        ocr_pages = _run_ocr_in_process(path)
        if ocr_pages and not _has_mojibake(ocr_pages):
            log.info("OCR erfolgreich: %s (%d Seiten)", path.name, len(ocr_pages))
            return ocr_pages
    except ExtractionError as e:
        log.warning("OCR fehlgeschlagen (%s): %s", path.name, e)
    except Exception as e:
        log.warning("OCR Fehler (%s): %s", path.name, e)

    raise ExtractionError(f"PDF konnte nicht gelesen werden: {path.name}")


CHUNK_SIZE    = 800   # Zeichen pro Chunk
CHUNK_OVERLAP = 120  # Überlapp zwischen benachbarten Chunks


def _split_long_text(page_number, text: str, max_len: int = 3000) -> list[dict]:
    """PDF-seitenweises Splitting — rekursiv bis alle Teile ≤ max_len Zeichen."""
    if len(text) <= max_len:
        return [{"page_number": page_number, "content": text}]
    mid = len(text) // 2
    split_at = text.rfind(" ", max(0, mid - 300), min(len(text), mid + 300))
    if split_at <= 0:
        split_at = mid
    left  = text[:split_at].strip()
    right = text[split_at:].strip()
    return (
        _split_long_text(page_number, left,  max_len) +
        _split_long_text(page_number, right, max_len)
    )


def split_text_into_chunks(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Teilt Text in überlappende Chunks auf. Trennt bevorzugt an Absatz- oder Wortgrenzen."""
    text = text.strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start  = 0
    while start < len(text):
        end = start + chunk_size
        if end >= len(text):
            chunks.append(text[start:].strip())
            break
        # Absatzgrenze bevorzugen
        split_at = text.rfind("\n\n", start, end)
        if split_at == -1 or split_at <= start:
            # Zeilenumbruch
            split_at = text.rfind("\n", start + chunk_size // 2, end)
        if split_at == -1 or split_at <= start:
            # Wortgrenze
            split_at = text.rfind(" ", start + chunk_size // 2, end)
        if split_at == -1 or split_at <= start:
            split_at = end
        chunk = text[start:split_at].strip()
        if chunk:
            chunks.append(chunk)
        # Fortschritt GARANTIEREN: normalerweise überlappt der nächste Chunk um `overlap`.
        # Aber wenn `split_at - overlap` nicht über den aktuellen Start hinauskommt, würde
        # der Start stehenbleiben → Endlosschleife. Das passiert z.B. wenn eine früh
        # gefundene Absatzgrenze (\n\n) im Fenster bleibt und rfind sie immer wieder
        # liefert (Dateien mit EINER Leerzeile + langem Fliesstext, z.B. Kamera-Logs).
        # split_at ist per Konstruktion immer > start → als Fallback vorwärts springen.
        next_start = split_at - overlap
        if next_start <= start:
            next_start = split_at
        start = next_start
        if start < 0:
            start = 0
    return [c for c in chunks if len(c) > 20]


def extract_chunks(path: Path) -> list[dict]:
    """Gibt [{page_number, chunk_index, content}] zurück.

    PDF: ein Chunk pro Seite (lange Seiten werden gesplittet).
    Andere Formate: Text wird in überlappende Chunks von ~800 Zeichen aufgeteilt.
    Alle Texte werden normalisiert (Ligaturen → reguläre Zeichen).
    """
    ext = path.suffix.lower()
    if ext == ".pdf":
        pages = extract_pdf_pages(path)
        chunks = []
        idx = 0
        for page_data in pages:
            for c in _split_long_text(page_data["page_number"], page_data["content"]):
                chunks.append({
                    "page_number": c["page_number"],
                    "chunk_index": idx,
                    "content":     normalize_text(c["content"]),
                })
                idx += 1
        return chunks
    else:
        # 30s Timeout für alle Nicht-PDF-Formate (DOCX, EML, RTF, TXT, …)
        use_alarm = _IN_WORKER_PROCESS
        if use_alarm:
            old_handler = _signal.signal(_signal.SIGALRM, _non_pdf_alarm)
            _signal.alarm(_NON_PDF_TIMEOUT)
        # Der Timeout umschliesst bewusst AUCH Normalisierung + Chunking — nicht nur das
        # Lesen. Sonst würde ein Hänger in split_text_into_chunks (kein interner Timeout)
        # erst vom 120s-SIGKILL des Scanners gestoppt.
        try:
            text, _ = extract(path)
            if not text:
                parts = []
            else:
                parts = split_text_into_chunks(normalize_text(text))
        except _NonPdfTimeout:
            log.warning("Extraktion Timeout (%ds): %s", _NON_PDF_TIMEOUT, path.name)
            raise ExtractionError(f"Timeout ({_NON_PDF_TIMEOUT}s): {path.name}")
        finally:
            if use_alarm:
                _signal.alarm(0)
                _signal.signal(_signal.SIGALRM, old_handler)
        return [
            {"page_number": None, "chunk_index": i, "content": part}
            for i, part in enumerate(parts)
        ]


# ── DOCX ──────────────────────────────────────────────────────────────────────

def extract_docx(path: Path) -> tuple[str, str]:
    try:
        import docx
    except ImportError:
        raise UnsupportedFormat("python-docx not installed")
    try:
        # READ-ONLY: NAS darf nie verändert werden — Document() öffnet nur lesend
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text), ""
    except Exception as exc:
        raise ExtractionError(str(exc)) from exc


# ── DOC (Word 97-2003) ────────────────────────────────────────────────────────

def extract_doc(path: Path) -> tuple[str, str]:
    """Nutzt macOS-integriertes textutil; kein extra Paket nötig."""
    import subprocess
    try:
        # READ-ONLY: NAS darf nie verändert werden — textutil liest nur
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            raise ExtractionError(result.stderr.strip() or "textutil fehlgeschlagen")
        return result.stdout, ""
    except FileNotFoundError:
        raise UnsupportedFormat("textutil nicht verfügbar (nur macOS)")
    except subprocess.TimeoutExpired:
        raise ExtractionError(f"textutil Timeout: {path}")


# ── XLSX ──────────────────────────────────────────────────────────────────────

_XLSX_MAX_MB      = 20       # Dateien grösser als 20 MB überspringen
_XLSX_MAX_SHEETS  = 10       # Maximal 10 Arbeitsblätter
_XLSX_MAX_ROWS    = 5_000    # Maximal 5000 Zeilen pro Blatt
_XLSX_TIMEOUT_SEC = 60       # Abbruch nach 60 Sekunden


def extract_xlsx(path: Path) -> tuple[str, str]:
    try:
        import openpyxl
    except ImportError:
        raise UnsupportedFormat("openpyxl not installed")

    # Grössencheck — sehr grosse XLSX-Dateien hängen openpyxl auf
    size_mb = path.stat().st_size / 1_048_576
    if size_mb > _XLSX_MAX_MB:
        raise ExtractionError(
            f"XLSX zu gross ({size_mb:.1f} MB > {_XLSX_MAX_MB} MB) — übersprungen"
        )

    # Extraktion in separatem Thread mit Timeout
    import threading
    result_box: list = []
    error_box:  list = []

    def _extract():
        try:
            # READ-ONLY: NAS darf nie verändert werden — read_only=True
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets[:_XLSX_MAX_SHEETS]:
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    line = "  ".join(
                        str(c)[:200] for c in row   # Zellinhalt auf 200 Zeichen begrenzen
                        if c is not None and str(c).strip()
                    )
                    if line:
                        parts.append(line)
                    row_count += 1
                    if row_count >= _XLSX_MAX_ROWS:
                        parts.append(f"[… gekürzt nach {_XLSX_MAX_ROWS} Zeilen]")
                        break
            wb.close()
            result_box.append(("\n".join(parts), ""))
        except Exception as exc:
            error_box.append(exc)

    t = threading.Thread(target=_extract, daemon=True)
    t.start()
    t.join(timeout=_XLSX_TIMEOUT_SEC)

    if t.is_alive():
        raise ExtractionError(
            f"XLSX Timeout nach {_XLSX_TIMEOUT_SEC}s — Datei zu komplex: {path.name}"
        )
    if error_box:
        raise ExtractionError(str(error_box[0]))
    return result_box[0] if result_box else ("", "")


# ── EML ───────────────────────────────────────────────────────────────────────

_MAX_EML_BYTES = 524288  # 512 KB — Header + Text, keine Anhang-Binärdaten in RAM


def extract_eml(path: Path) -> tuple[str, str]:
    import email
    try:
        # Nur erste 512 KB lesen: Anhänge (Base64-Blöcke) kommen danach — nie in RAM
        with open(path, "rb") as f:
            raw = f.read(_MAX_EML_BYTES)
        msg = email.message_from_bytes(raw)

        parts = []
        subject  = msg.get("subject", "")
        sender   = msg.get("from", "")
        to_addr  = msg.get("to", "")
        if subject:
            parts.append(f"Betreff: {subject}")
        if sender:
            parts.append(f"Von: {sender}")
        if to_addr:
            parts.append(f"An: {to_addr}")

        att_names = []
        if msg.is_multipart():
            for part in msg.walk():
                fname = part.get_filename()
                if fname:
                    att_names.append(fname)
                    continue  # Anhang: nur Dateiname, kein Inhalt
                disp = part.get("Content-Disposition", "") or ""
                if part.get_content_type() == "text/plain" and "attachment" not in disp:
                    try:
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or "utf-8"
                            parts.append(payload.decode(charset, errors="replace"))
                    except Exception:
                        pass
        else:
            try:
                payload = msg.get_payload(decode=True)
                if payload:
                    parts.append(payload.decode(msg.get_content_charset() or "utf-8", errors="replace"))
            except Exception:
                pass

        if att_names:
            parts.append("Anhänge: " + ", ".join(att_names))

        return "\n".join(parts), ""
    except Exception as exc:
        raise ExtractionError(str(exc)) from exc


# ── RTF ───────────────────────────────────────────────────────────────────────

def extract_rtf(path: Path) -> tuple[str, str]:
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        raise UnsupportedFormat("striprtf not installed")
    try:
        # READ-ONLY: NAS darf nie verändert werden
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                raw = path.read_text(encoding=enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ExtractionError(f"Cannot decode RTF: {path}")
        return rtf_to_text(raw), ""
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(str(exc)) from exc


# ── Registry ──────────────────────────────────────────────────────────────────

_REGISTRY: dict[str, callable] = {
    ".txt":  extract_txt,
    ".md":   extract_txt,
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".dotx": extract_docx,  # Word-Vorlage -- strukturell dasselbe OOXML-Format wie .docx
    ".doc":  extract_doc,
    ".xlsx": extract_xlsx,
    ".eml":  extract_eml,
    ".rtf":  extract_rtf,
}


def extract(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    fn = _REGISTRY.get(ext)
    if fn is None:
        raise UnsupportedFormat(f"No extractor for {ext}")
    return fn(path)
